from __future__ import annotations

from base64 import b64encode
import csv
import io
import json
import math
import os
import re
import shutil
import subprocess
import threading
import time
import unicodedata
import zipfile
from collections import Counter, defaultdict
from datetime import datetime, timedelta
from pathlib import Path
from typing import Any
from xml.etree import ElementTree as ET

from backend.app.core.platform_registry import (
    SYSTEM_FIELDS,
    find_domain as default_find_domain,
    registry_payload,
)
from backend.app.core.advanced_text_visualization import advanced_text_visualization_payload, clean_text
from backend.app.core.llm_client import configured as llm_configured, chat_completion


ROOT = Path(__file__).resolve().parents[3]
APP_DATA_ROOT = Path(os.environ.get("PLATFORM_DATA_ROOT") or os.environ.get("DATA_ROOT") or (ROOT / "data")).resolve()
UPLOAD_ROOT = Path(os.environ.get("UPLOAD_ROOT") or os.environ.get("PLATFORM_UPLOAD_ROOT") or (APP_DATA_ROOT / "uploads")).resolve()
BACKUP_ROOT = Path(os.environ.get("BACKUP_ROOT") or os.environ.get("PLATFORM_BACKUP_ROOT") or (APP_DATA_ROOT / "backups")).resolve()
DATABASE_BACKUP_ROOT = Path(os.environ.get("DATABASE_BACKUP_ROOT") or (BACKUP_ROOT / "database")).resolve()
FILE_BACKUP_ROOT = Path(os.environ.get("FILE_BACKUP_ROOT") or (BACKUP_ROOT / "files")).resolve()
STORE_PATH = APP_DATA_ROOT / "platform_store.json"
EXPORT_ROOT = APP_DATA_ROOT / "exports"
CHUNK_ROOT = APP_DATA_ROOT / "upload_chunks"
RETENTION_DAYS = int(os.environ.get("BACKUP_RETENTION_DAYS", "30") or "30")
OCR_PROVIDER = os.environ.get("PLATFORM_OCR_PROVIDER", "gpt").strip() or "gpt"
OCR_MODEL = os.environ.get("PLATFORM_OCR_MODEL", "").strip()
TEXT_CHUNK_SIZE = max(800, int(os.environ.get("PLATFORM_TEXT_CHUNK_SIZE", "2600") or "2600"))
SCHEDULER_STARTED = False
BUILTIN_TABLE_CACHE: dict[str, tuple[float, int, list[str], list[dict[str, Any]]]] = {}
IMAGE_EXTENSIONS = {"png", "jpg", "jpeg", "webp", "bmp", "tif", "tiff"}
SUPPORTED_UPLOAD_EXTENSIONS = {"xlsx", "xls", "xlsm", "csv", "tsv", "json", "pdf", "docx", *IMAGE_EXTENSIONS}
DOCUMENT_EXTENSIONS = {"pdf", "docx", *IMAGE_EXTENSIONS}
TABLE_EXTENSIONS = SUPPORTED_UPLOAD_EXTENSIONS - DOCUMENT_EXTENSIONS


FIELD_ALIASES = {
    "title": ["title", "name", "标题", "题名", "书名", "作品", "故事集标题", "规范故事名"],
    "author": ["author", "作者", "作家", "编者", "editor"],
    "translator": ["translator", "译者", "翻译", "译者/编者", "editor"],
    "publisher": ["publisher", "出版社", "出版机构"],
    "publish_year": ["year", "publish_year", "publicationyear", "出版年", "年份", "年代", "出版时间"],
    "country": ["country", "国家", "出版国家"],
    "city": ["city", "城市", "出版地", "出版城市"],
    "theme": ["theme", "topic", "主题", "主题词", "关键词", "类型"],
    "content": ["content", "text", "正文", "译文内容", "内容", "序跋文本", "简介", "摘要"],
    "source": ["source", "来源", "source region", "来源地", "省份"],
    "preface": ["preface", "序", "跋", "序跋", "序言"],
    "notes": ["notes", "备注", "说明", "note"],
}

STOPWORDS = {
    "的", "了", "和", "与", "及", "在", "为", "是", "对", "中", "中国", "一个", "the", "and", "of", "in", "to", "a", "an",
    "und", "der", "die", "das", "den", "dem", "des", "ein", "eine", "einer", "einen", "einem", "ist", "sind", "war", "waren",
    "nicht", "mit", "von", "aus", "auf", "für", "sich", "sie", "wie", "auch", "oder", "aber", "als", "bei", "nach", "noch",
    "nur", "über", "unter", "durch", "man", "wir", "ihr", "ihre", "ihren", "ihm", "ihn", "ich", "du", "er", "es", "doch",
    "dass", "hier", "diese", "dieser", "dieses", "seine", "seiner", "seines", "zur", "zum", "mir", "mich", "dir", "dich",
    "habe", "hast", "hat", "haben", "hatte", "hatten", "werde", "werden", "wird", "wurde", "wurden", "eines", "einer",
    "einem", "sehr", "dann", "wohl", "mittels", "lassen", "lässt", "laesst", "kann", "können", "koennen", "muss", "muß",
    "buch", "seite", "kapitel", "la", "le", "de", "des", "du",
    "les", "une", "est", "pour", "dans", "que", "qui", "sur", "par", "el", "los", "las", "una", "para", "por", "con",
    "pdf", "docx", "ocr", "txt", "文本", "页面", "提取", "测试", "混合", "解析", "识别", "上传", "文件", "内容",
}


def ensure_dirs() -> None:
    for path in [APP_DATA_ROOT, UPLOAD_ROOT, BACKUP_ROOT, DATABASE_BACKUP_ROOT, FILE_BACKUP_ROOT, EXPORT_ROOT, CHUNK_ROOT]:
        path.mkdir(parents=True, exist_ok=True)


def _now_iso() -> str:
    return datetime.now().astimezone().isoformat()


def _timestamp() -> str:
    return datetime.now().strftime("%Y%m%d_%H%M%S")


def _empty_store() -> dict[str, Any]:
    registry = registry_payload()
    return {
        "next_dataset_id": 1,
        "next_record_id": 1,
        "next_log_id": 1,
        "next_backup_id": 1,
        "domains": registry["domains"],
        "components": registry["components"],
        "deleted_submodule_ids": [],
        "datasets": [],
        "records": [],
        "field_mappings": [],
        "visualization_cache": [],
        "operation_logs": [],
        "backup_jobs": [],
        "system_config": {
            "email": {"enabled": True, "admin_notify": ""},
            "cache": {"ttl_minutes": 60, "auto_clear_on_upload": True},
            "backup": {
                "database_time": "02:00",
                "files_time": "03:00",
                "retention_days": RETENTION_DAYS,
                "database_backup_root": str(DATABASE_BACKUP_ROOT),
                "file_backup_root": str(FILE_BACKUP_ROOT),
                "upload_root": str(UPLOAD_ROOT),
            },
        },
    }


def _merge_registry_state(data: dict[str, Any]) -> dict[str, Any]:
    registry = registry_payload()
    stored_domains = {str(domain.get("id")): domain for domain in data.get("domains", []) if isinstance(domain, dict)}
    deleted_submodule_ids = _deleted_submodule_needles(data)
    merged_domains: list[dict[str, Any]] = []

    for default_domain in registry["domains"]:
        stored_domain = stored_domains.get(str(default_domain.get("id"))) or {}
        domain = {**default_domain, **{key: value for key, value in stored_domain.items() if key != "submodules"}}
        stored_submodules = {
            str(submodule.get("id")): submodule
            for submodule in stored_domain.get("submodules", [])
            if isinstance(submodule, dict)
        }
        merged_submodules = []
        for default_submodule in default_domain.get("submodules", []):
            submodule_id = str(default_submodule.get("id"))
            numeric_id = str(default_submodule.get("numericId"))
            if submodule_id in deleted_submodule_ids or numeric_id in deleted_submodule_ids:
                continue
            stored_submodule = stored_submodules.pop(submodule_id, {})
            merged_submodule = {**default_submodule, **stored_submodule}
            enabled = list(merged_submodule.get("enabled_components") or [])
            for required in ["advanced-text-visuals", "full-text"]:
                if required not in enabled:
                    enabled.append(required)
            if "检索" in str(merged_submodule.get("name") or "") and "word-distance" in enabled and "word-comparison" not in enabled:
                insert_at = enabled.index("word-distance") + 1
                enabled.insert(insert_at, "word-comparison")
            merged_submodule["enabled_components"] = enabled
            merged_submodules.append(merged_submodule)
        for submodule in stored_submodules.values():
            submodule_id = str(submodule.get("id"))
            numeric_id = str(submodule.get("numericId"))
            if submodule_id not in deleted_submodule_ids and numeric_id not in deleted_submodule_ids:
                enabled = list(submodule.get("enabled_components") or [])
                for required in ["advanced-text-visuals", "full-text"]:
                    if required not in enabled:
                        enabled.append(required)
                if "检索" in str(submodule.get("name") or "") and "word-distance" in enabled and "word-comparison" not in enabled:
                    insert_at = enabled.index("word-distance") + 1
                    enabled.insert(insert_at, "word-comparison")
                submodule["enabled_components"] = enabled
                merged_submodules.append(submodule)
        domain["submodules"] = sorted(
            merged_submodules,
            key=lambda item: (int(item.get("sort_order") or 999), str(item.get("name") or "")),
        )
        merged_domains.append(domain)

    known_domain_ids = {str(domain.get("id")) for domain in registry["domains"]}
    for domain_id, stored_domain in stored_domains.items():
        if domain_id not in known_domain_ids:
            merged_domains.append(stored_domain)

    data["domains"] = sorted(merged_domains, key=lambda item: int(item.get("sort_order") or 999))
    data["components"] = registry["components"]
    data.setdefault("deleted_submodule_ids", [])
    return data


def _deleted_submodule_needles(data: dict[str, Any]) -> set[str]:
    return {str(item) for item in data.get("deleted_submodule_ids", []) if item not in (None, "")}


def _matches_submodule(submodule: dict[str, Any], needle: str) -> bool:
    return str(submodule.get("id")) == needle or str(submodule.get("numericId")) == needle


def load_store() -> dict[str, Any]:
    ensure_dirs()
    if not STORE_PATH.exists():
        store = _empty_store()
        save_store(store)
        return store
    try:
        data = json.loads(STORE_PATH.read_text(encoding="utf-8"))
    except Exception:
        data = _empty_store()
    for key, default in _empty_store().items():
        data.setdefault(key, default)
    data = _merge_registry_state(data)
    return data


def all_submodules() -> list[dict[str, Any]]:
    return [submodule for domain in load_store().get("domains", []) for submodule in domain.get("submodules", [])]


def find_domain(domain_id_or_numeric: str | int) -> dict[str, Any] | None:
    needle = str(domain_id_or_numeric)
    for domain in load_store().get("domains", []):
        if str(domain.get("id")) == needle or str(domain.get("numericId")) == needle:
            return domain
    return default_find_domain(domain_id_or_numeric)


def find_submodule(sub_module_id: str | int) -> dict[str, Any] | None:
    needle = str(sub_module_id)
    store = load_store()
    if needle in _deleted_submodule_needles(store):
        return None
    for domain in store.get("domains", []):
        for submodule in domain.get("submodules", []):
            if _matches_submodule(submodule, needle):
                return submodule
    return None


def save_store(store: dict[str, Any]) -> None:
    ensure_dirs()
    tmp = STORE_PATH.with_suffix(".tmp")
    tmp.write_text(json.dumps(store, ensure_ascii=False, indent=2), encoding="utf-8")
    tmp.replace(STORE_PATH)


def record_operation(operation_type: str, content: str, user_id: int | None = None, ip_address: str = "", user_agent: str = "") -> dict[str, Any]:
    store = load_store()
    item = {
        "id": store.get("next_log_id", 1),
        "user_id": user_id,
        "operation_type": operation_type,
        "operation_content": content,
        "ip_address": ip_address,
        "user_agent": user_agent,
        "created_at": _now_iso(),
    }
    store["next_log_id"] = item["id"] + 1
    store.setdefault("operation_logs", []).append(item)
    save_store(store)
    return item


def platform_registry() -> dict[str, Any]:
    store = load_store()
    return {"domains": store["domains"], "components": store["components"], "systemFields": SYSTEM_FIELDS}


def _safe_filename(name: str) -> str:
    cleaned = re.sub(r'[<>:"/\\|?*\x00-\x1f]+', "_", name or "dataset")
    return cleaned.strip(" .") or "dataset"


def _extension(filename: str) -> str:
    return Path(filename or "").suffix.lower().lstrip(".")


def _read_docx_text(content: bytes) -> str:
    try:
        from docx import Document  # type: ignore

        document = Document(io.BytesIO(content))
        parts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text and paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                row_text = "\t".join(cell.text.strip() for cell in row.cells if cell.text and cell.text.strip())
                if row_text:
                    parts.append(row_text)
        return "\n".join(parts).strip()
    except Exception:
        pass

    try:
        with zipfile.ZipFile(io.BytesIO(content)) as archive:
            xml = archive.read("word/document.xml")
    except Exception:
        return ""
    root = ET.fromstring(xml)
    ns = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    }
    paragraphs = []
    for paragraph in root.findall(".//w:p", ns):
        text = "".join(node.text or "" for node in paragraph.findall(".//w:t", ns)).strip()
        if text:
            paragraphs.append(text)
    return "\n".join(paragraphs).strip()


def _read_pdf_pages(content: bytes) -> list[str]:
    try:
        import fitz  # type: ignore

        document = fitz.open(stream=content, filetype="pdf")
        parts: list[str] = []
        for page in document:
            try:
                parts.append((page.get_text("text") or "").strip())
            except Exception:
                parts.append("")
        if any(part.strip() for part in parts):
            return parts
    except Exception:
        pass

    for module_name, reader_name in [("pypdf", "PdfReader"), ("PyPDF2", "PdfReader")]:
        try:
            module = __import__(module_name, fromlist=[reader_name])
            reader = getattr(module, reader_name)(io.BytesIO(content))
            pages = []
            for page in reader.pages:
                try:
                    pages.append((page.extract_text() or "").strip())
                except Exception:
                    pages.append("")
            if any(part.strip() for part in pages):
                return pages
        except Exception:
            continue
    return []


def _read_pdf_text(content: bytes) -> str:
    return "\n\n".join(
        f"[第 {index} 页]\n{text.strip()}"
        for index, text in enumerate(_read_pdf_pages(content), start=1)
        if text.strip()
    ).strip()


def _pdf_page_images(content: bytes) -> list[bytes]:
    images: list[bytes] = []
    try:
        import fitz  # type: ignore

        document = fitz.open(stream=content, filetype="pdf")
        for page in document:
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
            images.append(pix.tobytes("png"))
        return images
    except Exception:
        return images


def _extract_image_text_with_llm(image_bytes: bytes, filename: str) -> str:
    if not llm_configured(OCR_PROVIDER, OCR_MODEL):
        return ""
    mime = "image/png"
    ext = _extension(filename)
    if ext in {"jpg", "jpeg"}:
        mime = "image/jpeg"
    elif ext == "webp":
        mime = "image/webp"
    elif ext in {"tif", "tiff"}:
        mime = "image/tiff"
    data_url = f"data:{mime};base64,{b64encode(image_bytes).decode('ascii')}"
    messages = [
        {"role": "system", "content": "你是一名OCR识别引擎。请只输出图片中的可见文字，尽量保留段落、表格与换行，不要解释。"},
        {"role": "user", "content": [{"type": "text", "text": "请识别这张图片中的文字内容。"}, {"type": "image_url", "image_url": {"url": data_url}}]},
    ]
    return chat_completion(messages, model=OCR_MODEL or "general", provider=OCR_PROVIDER, timeout=120).strip()


def _extract_pdf_text_with_llm_ocr(content: bytes, filename: str) -> tuple[str, int]:
    if not llm_configured(OCR_PROVIDER, OCR_MODEL):
        return "", 0
    parts = []
    processed = 0
    for index, image_bytes in enumerate(_pdf_page_images(content), start=1):
        try:
            text = _extract_image_text_with_llm(image_bytes, f"{Path(filename).stem}-page-{index}.png")
        except Exception:
            text = ""
        processed += 1
        if text.strip():
            parts.append(f"[第 {index} 页]\n{text.strip()}")
    return "\n\n".join(parts).strip(), processed


def _extract_uploaded_text(content: bytes, filename: str, force_ocr: bool = False) -> tuple[str, str, dict[str, Any]]:
    ext = _extension(filename)
    if ext == "docx":
        text = _read_docx_text(content)
        return text, "docx", {"page_count": 0, "ocr_pages": 0, "direct_text": True}
    if ext == "pdf":
        pages = _read_pdf_pages(content)
        direct_text = "\n\n".join(
            f"[第 {index} 页]\n{text.strip()}"
            for index, text in enumerate(pages, start=1)
            if text.strip()
        ).strip()
        if force_ocr:
            ocr_text, ocr_pages = _extract_pdf_text_with_llm_ocr(content, filename)
            if ocr_text.strip():
                return ocr_text, "pdf-ocr", {"page_count": max(len(pages), ocr_pages), "ocr_pages": ocr_pages, "direct_text": False}
            if direct_text:
                return direct_text, "pdf", {"page_count": len(pages), "ocr_pages": 0, "direct_text": True, "ocr_fallback": True}
        text = _read_pdf_text(content)
        if text.strip():
            return text, "pdf", {"page_count": len(pages), "ocr_pages": 0, "direct_text": True}
        ocr_text, ocr_pages = _extract_pdf_text_with_llm_ocr(content, filename)
        return ocr_text, "pdf-ocr", {"page_count": max(len(pages), ocr_pages), "ocr_pages": ocr_pages, "direct_text": False}
    if ext in IMAGE_EXTENSIONS:
        text = _extract_image_text_with_llm(content, filename)
        return text, "ocr" if text.strip() else "image", {"page_count": 1, "ocr_pages": 1 if text.strip() else 0, "direct_text": False}
    return "", ext, {"page_count": 0, "ocr_pages": 0, "direct_text": False}


def _split_text_chunks(text: str, size: int = TEXT_CHUNK_SIZE) -> list[str]:
    clean = re.sub(r"\r\n?", "\n", text or "").strip()
    if not clean:
        return []
    paragraphs = [part.strip() for part in re.split(r"\n{2,}", clean) if part.strip()]
    chunks: list[str] = []
    buffer = ""
    for paragraph in paragraphs:
        if len(paragraph) > size:
            if buffer.strip():
                chunks.append(buffer.strip())
                buffer = ""
            for index in range(0, len(paragraph), size):
                chunk = paragraph[index:index + size].strip()
                if chunk:
                    chunks.append(chunk)
            continue
        candidate = paragraph if not buffer else f"{buffer}\n\n{paragraph}"
        if len(candidate) <= size:
            buffer = candidate
        else:
            if buffer.strip():
                chunks.append(buffer.strip())
            buffer = paragraph
    if buffer.strip():
        chunks.append(buffer.strip())
    return chunks


def _text_rows_from_document(text: str, source_name: str, kind: str) -> tuple[list[str], list[dict[str, Any]]]:
    chunks = _split_text_chunks(text)
    if not chunks and text.strip():
        chunks = [text.strip()]
    headers = ["title", "content", "source", "notes"]
    source_title = Path(source_name).stem
    rows = [
        {
            "title": f"{source_title} · 文本片段 {index}",
            "content": chunk,
            "source": source_name,
            "notes": f"{kind} · 第 {index}/{len(chunks)} 个文本片段",
        }
        for index, chunk in enumerate(chunks, start=1)
    ]
    if not rows and text.strip():
        rows = [{"title": source_title, "content": text.strip(), "source": source_name, "notes": kind}]
    return headers, rows


def _document_analysis(text: str, chunks: list[str], extraction: dict[str, Any]) -> dict[str, Any]:
    clean = re.sub(r"\r\n?", "\n", text or "").strip()
    paragraphs = [part.strip() for part in re.split(r"\n{2,}", clean) if part.strip()]
    sentences = [part.strip() for part in re.split(r"(?<=[。！？.!?])\s+", clean) if part.strip()]
    tokens = _tokenize(clean)
    counts = Counter(tokens)
    document_frequency: Counter[str] = Counter()
    for chunk in chunks:
        document_frequency.update(set(_tokenize(chunk)))
    chunk_count = max(1, len(chunks))
    scored = []
    for word, count in counts.items():
        score = count * (1.0 + max(0.0, __import__("math").log((chunk_count + 1) / (document_frequency[word] + 1))))
        scored.append((score, word, count))
    scored.sort(reverse=True)
    top_keywords = [{"word": word, "count": count, "score": round(score, 3)} for score, word, count in scored[:48]]

    cooccurrence: Counter[tuple[str, str]] = Counter()
    important = {item["word"] for item in top_keywords[:24]}
    for chunk in chunks:
        window = [token for token in _tokenize(chunk) if token in important]
        for index, left in enumerate(window):
            for right in window[index + 1:index + 7]:
                if left != right:
                    cooccurrence[tuple(sorted((left, right)))] += 1

    chapter_pattern = re.compile(
        r"^(?:第[一二三四五六七八九十百千万0-9]+[章节卷部篇]|chapter\s+\d+|contents|目录|引言|前言|绪论|结语|参考文献).{0,80}$",
        re.I,
    )
    headings = []
    for line in clean.splitlines():
        candidate = line.strip()
        if candidate and len(candidate) <= 90 and chapter_pattern.match(candidate):
            headings.append(candidate)
        if len(headings) >= 80:
            break

    chinese_chars = len(re.findall(r"[\u4e00-\u9fff]", clean))
    latin_words = len(re.findall(r"\b[A-Za-z][A-Za-z'-]*\b", clean))
    language = "中文为主" if chinese_chars >= latin_words * 2 else "外文或混合文本"
    return {
        "char_count": len(clean),
        "page_count": int(extraction.get("page_count") or 0),
        "ocr_pages": int(extraction.get("ocr_pages") or 0),
        "paragraph_count": len(paragraphs),
        "sentence_count": len(sentences),
        "chunk_count": len(chunks),
        "token_count": len(tokens),
        "unique_term_count": len(counts),
        "reading_minutes": max(1, round((chinese_chars + latin_words * 5) / 420)),
        "language": language,
        "top_keywords": top_keywords,
        "cooccurrence": [
            {"source": source, "target": target, "count": count}
            for (source, target), count in cooccurrence.most_common(40)
        ],
        "headings": headings,
    }


def parse_field_mappings(raw: Any, headers: list[str] | None = None) -> dict[str, str]:
    if raw is None or raw == "":
        raw = {}
    if isinstance(raw, str):
        try:
            raw = json.loads(raw)
        except Exception:
            raw = {}
    result: dict[str, str] = {}
    if isinstance(raw, list):
        for item in raw:
            if not isinstance(item, dict):
                continue
            file_column = str(item.get("file_column") or item.get("fileColumn") or item.get("column") or "").strip()
            system_field = str(item.get("system_field") or item.get("systemField") or item.get("field") or "").strip()
            if file_column and system_field in SYSTEM_FIELDS:
                result[file_column] = system_field
    elif isinstance(raw, dict):
        for key, value in raw.items():
            key_s = str(key).strip()
            value_s = str(value).strip()
            if value_s in SYSTEM_FIELDS:
                result[key_s] = value_s
            elif key_s in SYSTEM_FIELDS:
                result[value_s] = key_s
    if not result and headers:
        normalized_headers = {header: re.sub(r"\s+", "", str(header).lower()) for header in headers}
        for system_field, aliases in FIELD_ALIASES.items():
            for header, normalized in normalized_headers.items():
                if any(re.sub(r"\s+", "", alias.lower()) in normalized for alias in aliases):
                    result[header] = system_field
                    break
    return result


def _rows_from_xlsx_bytes(content: bytes) -> tuple[list[str], list[dict[str, Any]]]:
    try:
        import openpyxl  # type: ignore

        workbook = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)
        sheet = workbook[workbook.sheetnames[0]]
        iterator = sheet.iter_rows(values_only=True)
        headers = [str(value or f"field_{index + 1}").strip() for index, value in enumerate(next(iterator, []))]
        rows = []
        for raw in iterator:
            if not raw or not any(value not in (None, "") for value in raw):
                continue
            rows.append({headers[index] if index < len(headers) else f"field_{index + 1}": value for index, value in enumerate(raw)})
        return headers, rows
    except Exception:
        pass

    with zipfile.ZipFile(io.BytesIO(content)) as archive:
        shared = []
        if "xl/sharedStrings.xml" in archive.namelist():
            root = ET.fromstring(archive.read("xl/sharedStrings.xml"))
            ns = {"a": "http://schemas.openxmlformats.org/spreadsheetml/2006/main"}
            for item in root.findall("a:si", ns):
                shared.append("".join(node.text or "" for node in item.findall(".//a:t", ns)))
        sheet_name = next(name for name in archive.namelist() if name.startswith("xl/worksheets/sheet"))
        root = ET.fromstring(archive.read(sheet_name))
        ns = {"a": "http://schemas.openxmlformats.org/spreadsheetml/2006/main"}
        matrix = []
        for row in root.findall(".//a:row", ns):
            values = []
            for cell in row.findall("a:c", ns):
                value = cell.find("a:v", ns)
                text = value.text if value is not None else ""
                if cell.get("t") == "s" and text.isdigit() and int(text) < len(shared):
                    text = shared[int(text)]
                values.append(text)
            matrix.append(values)
        headers = [str(value or f"field_{index + 1}").strip() for index, value in enumerate(matrix[0] if matrix else [])]
        return headers, [{headers[index] if index < len(headers) else f"field_{index + 1}": value for index, value in enumerate(row)} for row in matrix[1:]]


def parse_tabular_file(content: bytes, filename: str) -> tuple[list[str], list[dict[str, Any]]]:
    ext = _extension(filename)
    if ext in {"xlsx", "xlsm", "xls"}:
        return _rows_from_xlsx_bytes(content)
    text = content.decode("utf-8-sig", errors="replace")
    if ext == "json":
        payload = json.loads(text)
        rows = payload if isinstance(payload, list) else payload.get("items") or payload.get("rows") or payload.get("data") or []
        if rows and isinstance(rows[0], list):
            headers = [str(value or f"field_{index + 1}") for index, value in enumerate(rows[0])]
            return headers, [{headers[index] if index < len(headers) else f"field_{index + 1}": value for index, value in enumerate(row)} for row in rows[1:]]
        if rows and isinstance(rows[0], dict):
            headers = list(rows[0].keys())
            return headers, rows
        return [], []
    first_line = text.splitlines()[0] if text.splitlines() else ""
    delimiter = "\t" if ext == "tsv" or "\t" in first_line else ","
    reader = csv.DictReader(io.StringIO(text), delimiter=delimiter)
    headers = list(reader.fieldnames or [])
    return headers, [dict(row) for row in reader if any(str(value or "").strip() for value in row.values())]


def _record_values(row: dict[str, Any], headers: list[str], mappings: dict[str, str]) -> dict[str, Any]:
    fields = {f"field_{index + 1}": "" for index in range(20)}
    for index, header in enumerate(headers[:20]):
        fields[f"field_{index + 1}"] = "" if row.get(header) is None else str(row.get(header))
    system = {field: "" for field in SYSTEM_FIELDS}
    for header, system_field in mappings.items():
        if system_field in system:
            system[system_field] = "" if row.get(header) is None else str(row.get(header))
    if not system["title"] and headers:
        system["title"] = str(row.get(headers[0]) or "")
    if not system["content"]:
        system["content"] = " ".join(str(row.get(header) or "") for header in headers)
    return {**fields, "system": system, "raw": {header: row.get(header) for header in headers}}


def _cell(row: dict[str, Any], *keys: str) -> str:
    for key in keys:
        value = row.get(key)
        if value is not None and str(value).strip():
            return str(value).strip()
    return ""


def _builtin_specs_for_submodule(sub_module_id: int | str) -> list[dict[str, Any]]:
    store = load_store()
    if str(sub_module_id) in _deleted_submodule_needles(store):
        return []
    submodule = find_submodule(sub_module_id)
    sub_id = str(submodule.get("id") if submodule else sub_module_id)
    if sub_id in {"stories-overview", "stories-german-story-atlas"}:
        return [{
            "id": "builtin-story-collections",
            "filename": "中国故事集总表_知识库.xlsx",
            "name": "中国故事集总表",
            "kind": "story_collection",
        }]
    if sub_id == "stories-preface-atlas":
        return [{
            "id": "builtin-story-prefaces",
            "filename": "中国故事集_序跋.xlsx",
            "name": "中国故事集序跋",
            "kind": "story_preface",
        }]
    if sub_id == "stories-child-story-atlas":
        return [{
            "id": "builtin-story-children",
            "filename": "中国故事集_子故事（3533篇）.xlsx",
            "name": "中国故事集子故事",
            "kind": "story_child",
        }]
    if sub_id == "stories-wilhelm":
        return [
            {
                "id": "builtin-wilhelm-stories",
                "filename": "中国民间童话.xlsx",
                "name": "卫礼贤《中国民间童话》单篇译文",
                "kind": "wilhelm_story",
            },
            {
                "id": "builtin-wilhelm-publications",
                "filename": "地图_中国民间童话.xlsx",
                "name": "卫礼贤《中国民间童话》再版及传播",
                "kind": "wilhelm_publication",
            },
        ]
    return []


def _builtin_rows(path: Path) -> tuple[list[str], list[dict[str, Any]]]:
    if not path.exists():
        return [], []
    stat = path.stat()
    cache_key = str(path.resolve())
    cached = BUILTIN_TABLE_CACHE.get(cache_key)
    if cached and cached[0] == stat.st_mtime and cached[1] == stat.st_size:
        return cached[2], cached[3]
    headers, rows = parse_tabular_file(path.read_bytes(), path.name)
    cleaned_rows = [row for row in rows if any(str(value or "").strip() for value in row.values())]
    BUILTIN_TABLE_CACHE[cache_key] = (stat.st_mtime, stat.st_size, headers, cleaned_rows)
    return headers, cleaned_rows


def _builtin_system_fields(kind: str, row: dict[str, Any]) -> dict[str, str]:
    if kind == "story_collection":
        return {
            "title": _cell(row, "故事集标题", "title"),
            "author": "",
            "translator": _cell(row, "译者/编者"),
            "publisher": _cell(row, "出版社", "publisher"),
            "publish_year": _cell(row, "年份", "year"),
            "country": _cell(row, "国家", "country"),
            "city": _cell(row, "城市", "city"),
            "theme": _cell(row, "故事集标题（中文）"),
            "content": " ".join(filter(None, [_cell(row, "故事集标题"), _cell(row, "故事集标题（中文）"), _cell(row, "译者/编者身份")])),
            "source": "",
            "preface": _cell(row, "序跋作者"),
            "notes": _cell(row, "译者/编者身份"),
        }
    if kind == "story_preface":
        return {
            "title": _cell(row, "故事集标题"),
            "author": _cell(row, "序跋作者"),
            "translator": _cell(row, "序跋作者"),
            "publisher": "",
            "publish_year": _cell(row, "年份"),
            "country": "",
            "city": "",
            "theme": _cell(row, "序跋类型"),
            "content": _cell(row, "序跋文本", "故事集介绍"),
            "source": _cell(row, "故事集介绍"),
            "preface": _cell(row, "序跋文本"),
            "notes": _cell(row, "序跋类型"),
        }
    if kind == "story_child":
        return {
            "title": _cell(row, "子故事标题", "规范故事名"),
            "author": "",
            "translator": "",
            "publisher": "",
            "publish_year": _cell(row, "年份"),
            "country": "",
            "city": "",
            "theme": _cell(row, "规范故事名"),
            "content": " ".join(filter(None, [_cell(row, "子故事标题"), _cell(row, "规范故事名"), _cell(row, "故事民族来源"), _cell(row, "故事集标题")])),
            "source": _cell(row, "故事民族来源"),
            "preface": "",
            "notes": _cell(row, "故事集标题"),
        }
    if kind == "wilhelm_story":
        return {
            "title": _cell(row, "单篇译文故事名"),
            "author": "卫礼贤",
            "translator": "卫礼贤",
            "publisher": "Eugen Diederichs Verlag",
            "publish_year": "1914",
            "country": "Germany",
            "city": "Jena",
            "theme": _cell(row, "卫礼贤的分类"),
            "content": _cell(row, "译文内容"),
            "source": _cell(row, "故事来源"),
            "preface": "",
            "notes": _cell(row, "卫礼贤的分类"),
        }
    if kind == "wilhelm_publication":
        return {
            "title": _cell(row, "titel"),
            "author": "卫礼贤",
            "translator": "卫礼贤",
            "publisher": _cell(row, "publisher"),
            "publish_year": _cell(row, "year"),
            "country": _cell(row, "country"),
            "city": _cell(row, "city"),
            "theme": _cell(row, "全/选/改编"),
            "content": " ".join(filter(None, [_cell(row, "titel"), _cell(row, "publisher"), _cell(row, "city"), _cell(row, "year")])),
            "source": _cell(row, "city"),
            "preface": "",
            "notes": _cell(row, "全/选/改编"),
        }
    return {field: "" for field in SYSTEM_FIELDS}


def _builtin_extra_fields(kind: str, row: dict[str, Any]) -> dict[str, str]:
    if kind == "story_collection":
        return {
            "title": _cell(row, "故事集标题", "title"),
            "translator": _cell(row, "译者/编者"),
            "publisher": _cell(row, "出版社", "publisher"),
            "publish_year": _cell(row, "年份", "year"),
            "country": _cell(row, "国家", "country"),
            "city": _cell(row, "城市", "city"),
            "language": "德语",
            "document_type": "故事集",
            "collection": _cell(row, "故事集标题"),
            "reprint_count": "",
            "preface_author": _cell(row, "序跋作者"),
            "keywords": _cell(row, "译者/编者身份"),
            "motif_type": "",
            "source_place": "",
        }
    if kind == "story_preface":
        return {
            "title": _cell(row, "故事集标题"),
            "translator": _cell(row, "序跋作者"),
            "publisher": "",
            "publish_year": _cell(row, "年份"),
            "country": "",
            "city": "",
            "language": "德语",
            "document_type": "序跋",
            "collection": _cell(row, "故事集标题"),
            "preface_author": _cell(row, "序跋作者"),
            "keywords": _cell(row, "序跋类型"),
            "motif_type": _cell(row, "序跋类型"),
            "source_place": "",
        }
    if kind == "story_child":
        return {
            "title": _cell(row, "子故事标题", "规范故事名"),
            "translator": "",
            "publisher": "",
            "publish_year": _cell(row, "年份"),
            "country": "",
            "city": "",
            "language": "德语",
            "document_type": "子故事",
            "collection": _cell(row, "故事集标题"),
            "preface_author": "",
            "keywords": _cell(row, "规范故事名"),
            "motif_type": _cell(row, "规范故事名"),
            "source_place": _cell(row, "故事民族来源"),
        }
    if kind == "wilhelm_story":
        return {
            "title": _cell(row, "单篇译文故事名"),
            "translator": "卫礼贤",
            "publisher": "Eugen Diederichs Verlag",
            "publish_year": "1914",
            "country": "Germany",
            "city": "Jena",
            "language": "德语",
            "document_type": "单篇译文",
            "collection": "Chinesische Volksmärchen",
            "reprint_count": "",
            "preface_author": "卫礼贤",
            "keywords": _cell(row, "卫礼贤的分类"),
            "motif_type": _cell(row, "卫礼贤的分类"),
            "source_place": _cell(row, "故事来源"),
        }
    if kind == "wilhelm_publication":
        return {
            "title": _cell(row, "titel"),
            "translator": "卫礼贤",
            "publisher": _cell(row, "publisher"),
            "publish_year": _cell(row, "year"),
            "country": _cell(row, "country"),
            "city": _cell(row, "city"),
            "language": _cell(row, "语种") or "德语",
            "document_type": _cell(row, "全/选/改编") or "再版传播",
            "collection": _cell(row, "titel"),
            "reprint_count": "",
            "preface_author": "卫礼贤",
            "keywords": _cell(row, "全/选/改编"),
            "motif_type": _cell(row, "全/选/改编"),
            "source_place": _cell(row, "city"),
        }
    return {}


def _builtin_records_for_submodule(sub_module_id: int | str) -> list[dict[str, Any]]:
    records: list[dict[str, Any]] = []
    for spec in _builtin_specs_for_submodule(sub_module_id):
        path = APP_DATA_ROOT / spec["filename"]
        headers, rows = _builtin_rows(path)
        for index, row in enumerate(rows, start=1):
            fields = {f"field_{field_index + 1}": "" for field_index in range(20)}
            for field_index, header in enumerate(headers[:20]):
                fields[f"field_{field_index + 1}"] = str(row.get(header) or "")
            system = {field: "" for field in SYSTEM_FIELDS}
            system.update(_builtin_system_fields(spec["kind"], row))
            records.append({
                "id": f"{spec['id']}-{index}",
                "dataset_id": spec["id"],
                **fields,
                **_builtin_extra_fields(spec["kind"], row),
                "system": system,
                "raw": {header: row.get(header) for header in headers},
                "created_at": "",
                "updated_at": "",
                "builtin": True,
            })
    return records


def _builtin_dataset_metas(sub_module_id: int | str) -> list[dict[str, Any]]:
    metas = []
    submodule = find_submodule(sub_module_id)
    for spec in _builtin_specs_for_submodule(sub_module_id):
        path = APP_DATA_ROOT / spec["filename"]
        headers, rows = _builtin_rows(path)
        metas.append({
            "id": spec["id"],
            "sub_module_id": submodule.get("numericId") if submodule else sub_module_id,
            "subModuleId": submodule.get("id") if submodule else str(sub_module_id),
            "domainId": submodule.get("domainId") if submodule else "",
            "name": spec["name"],
            "file_name": spec["filename"],
            "file_path": str(path),
            "file_size": path.stat().st_size if path.exists() else 0,
            "file_type": "xlsx",
            "record_count": len(rows),
            "field_count": len(headers),
            "status": "completed" if path.exists() else "missing",
            "builtin": True,
        })
    return metas


def _dataset_upload_dir(dataset_id: int) -> Path:
    now = datetime.now()
    return UPLOAD_ROOT / f"{now.year:04d}" / f"{now.month:02d}" / f"{now.day:02d}" / str(dataset_id)


def upload_dataset_file(
    sub_module_id: str | int,
    file_bytes: bytes,
    filename: str,
    field_mappings: Any = None,
    dataset_name: str = "",
    description: str = "",
    affected_pages: Any = None,
    key_fields: Any = None,
    uploaded_by: int | None = None,
    dataset_kind: str = "",
    force_ocr: bool = False,
) -> dict[str, Any]:
    store = load_store()
    submodule = find_submodule(sub_module_id)
    if not submodule:
        raise ValueError("sub_module_id does not match a configured submodule.")
    dataset_id = int(store.get("next_dataset_id", 1))
    store["next_dataset_id"] = dataset_id + 1
    ext = _extension(filename)
    if ext not in SUPPORTED_UPLOAD_EXTENSIONS:
        raise ValueError("Unsupported file type. Use xlsx, csv, json, pdf, docx, or images.")
    resolved_kind = dataset_kind if dataset_kind in {"table", "document"} else ("document" if ext in DOCUMENT_EXTENSIONS else "table")
    if resolved_kind == "table" and ext not in TABLE_EXTENSIONS:
        raise ValueError("Table maintenance only accepts xlsx, xls, csv, tsv, or json files.")
    if resolved_kind == "document" and ext not in DOCUMENT_EXTENSIONS:
        raise ValueError("Document maintenance only accepts pdf, docx, or image files.")
    upload_dir = _dataset_upload_dir(dataset_id)
    upload_dir.mkdir(parents=True, exist_ok=True)
    safe_name = _safe_filename(filename)
    file_path = upload_dir / safe_name
    file_path.write_bytes(file_bytes)

    dataset = {
        "id": dataset_id,
        "sub_module_id": submodule["numericId"],
        "subModuleId": submodule["id"],
        "domainId": submodule["domainId"],
        "name": dataset_name or Path(filename).stem,
        "description": description,
        "affected_pages": json.loads(affected_pages) if isinstance(affected_pages, str) and affected_pages.strip().startswith("[") else (affected_pages or []),
        "key_fields": json.loads(key_fields) if isinstance(key_fields, str) and key_fields.strip().startswith("[") else (key_fields or []),
        "file_name": filename,
        "file_path": str(file_path),
        "file_size": len(file_bytes),
        "file_type": "xlsx" if ext in {"xls", "xlsm"} else ext,
        "dataset_kind": resolved_kind,
        "record_count": 0,
        "field_count": 0,
        "uploaded_by": uploaded_by,
        "status": "parsing",
        "error_message": "",
        "force_ocr": bool(force_ocr),
        "created_at": _now_iso(),
        "updated_at": _now_iso(),
    }
    store["datasets"].append(dataset)
    save_store(store)

    try:
        extracted_text, extraction_kind, extraction_meta = _extract_uploaded_text(file_bytes, filename, force_ocr=force_ocr)
        if resolved_kind == "document":
            headers, rows = _text_rows_from_document(extracted_text, filename, extraction_kind)
        else:
            headers, rows = parse_tabular_file(file_bytes, filename)
        mappings = parse_field_mappings(field_mappings, headers)
        store = load_store()
        next_record_id = int(store.get("next_record_id", 1))
        records = []
        for row in rows:
            values = _record_values(row, headers, mappings)
            records.append({
                "id": next_record_id,
                "dataset_id": dataset_id,
                **{f"field_{index + 1}": values[f"field_{index + 1}"] for index in range(20)},
                "system": values["system"],
                "raw": values["raw"],
                "text_kind": extraction_kind if extraction_kind in {"docx", "pdf", "pdf-ocr", "ocr", "image"} else "upload",
                "created_at": _now_iso(),
                "updated_at": _now_iso(),
            })
            next_record_id += 1
        store["next_record_id"] = next_record_id
        store.setdefault("records", []).extend(records)
        store["field_mappings"] = [item for item in store.get("field_mappings", []) if item.get("dataset_id") != dataset_id]
        store["field_mappings"].extend([
            {"dataset_id": dataset_id, "file_column": column, "system_field": field, "created_at": _now_iso(), "updated_at": _now_iso()}
            for column, field in mappings.items()
        ])
        for item in store["datasets"]:
            if item["id"] == dataset_id:
                item["record_count"] = len(records)
                item["field_count"] = len(headers)
                if resolved_kind == "document":
                    chunks = [str(row.get("content") or "") for row in rows if str(row.get("content") or "").strip()]
                    item["text_kind"] = extraction_kind if extracted_text.strip() else "pending_ocr"
                    item["preview_text"] = extracted_text[:6000]
                    item["analysis"] = _document_analysis(extracted_text, chunks, extraction_meta)
                    item["extraction"] = extraction_meta
                item["status"] = "completed" if rows or resolved_kind == "table" else "pending_ocr"
                item["updated_at"] = _now_iso()
                break
        store["visualization_cache"] = [
            item for item in store.get("visualization_cache", [])
            if str(item.get("sub_module_id")) not in {str(submodule["numericId"]), str(submodule["id"])}
        ]
        save_store(store)
        record_operation("dataset_upload", f"Uploaded dataset {dataset_id}: {filename}", uploaded_by)
    except Exception as error:
        store = load_store()
        for item in store["datasets"]:
            if item["id"] == dataset_id:
                item["status"] = "failed"
                item["error_message"] = str(error)
                item["updated_at"] = _now_iso()
                break
        save_store(store)
        if resolved_kind == "document":
            record_operation("document_parse_failed", f"Document {dataset_id} parse failed: {error}", uploaded_by)
            return get_dataset(dataset_id)
        raise
    return get_dataset(dataset_id)


def save_upload_chunk(upload_id: str, index: int, total: int, chunk: bytes, filename: str, sub_module_id: str | int, field_mappings: Any = None) -> dict[str, Any]:
    chunk_dir = CHUNK_ROOT / _safe_filename(upload_id)
    chunk_dir.mkdir(parents=True, exist_ok=True)
    (chunk_dir / f"{index:06d}.part").write_bytes(chunk)
    meta = {"filename": filename, "total": total, "sub_module_id": sub_module_id, "field_mappings": field_mappings}
    (chunk_dir / "meta.json").write_text(json.dumps(meta, ensure_ascii=False), encoding="utf-8")
    received = len(list(chunk_dir.glob("*.part")))
    if received < total:
        return {"upload_id": upload_id, "received": received, "total": total, "progress": received / max(1, total), "completed": False}
    content = b"".join((chunk_dir / f"{part:06d}.part").read_bytes() for part in range(total))
    dataset = upload_dataset_file(sub_module_id=sub_module_id, file_bytes=content, filename=filename, field_mappings=field_mappings)
    shutil.rmtree(chunk_dir, ignore_errors=True)
    return {"upload_id": upload_id, "received": total, "total": total, "progress": 1, "completed": True, "dataset": dataset}


def get_dataset(dataset_id: int | str) -> dict[str, Any]:
    store = load_store()
    for dataset in store.get("datasets", []):
        if str(dataset.get("id")) == str(dataset_id):
            return dict(dataset)
    raise KeyError("Dataset not found.")


def list_platform_datasets(sub_module_id: str | int | None = None) -> list[dict[str, Any]]:
    store = load_store()
    datasets = list(store.get("datasets", []))
    if sub_module_id not in (None, ""):
        if str(sub_module_id) in _deleted_submodule_needles(store):
            return []
        submodule = find_submodule(sub_module_id)
        if not submodule:
            return []
        needles = {str(sub_module_id)}
        needles.update({str(submodule["id"]), str(submodule["numericId"])})
        datasets = [dataset for dataset in datasets if str(dataset.get("sub_module_id")) in needles or str(dataset.get("subModuleId")) in needles]
    return datasets


def list_documents(sub_module_id: str | int | None = None) -> list[dict[str, Any]]:
    return [dataset for dataset in list_platform_datasets(sub_module_id) if dataset.get("dataset_kind") == "document"]


def document_detail(dataset_id: int | str, include_text: bool = True) -> dict[str, Any]:
    dataset = get_dataset(dataset_id)
    if dataset.get("dataset_kind") != "document":
        raise ValueError("The selected dataset is not a document.")
    records = _records_for_dataset(dataset_id)
    chunks = [
        {
            "id": record.get("id"),
            "title": _record_field(record, "title"),
            "content": _record_field(record, "content"),
            "notes": _record_field(record, "notes"),
            "text_kind": record.get("text_kind"),
        }
        for record in records
    ]
    payload = {"dataset": dataset, "chunks": chunks, "analysis": dataset.get("analysis") or {}}
    if include_text:
        payload["text"] = "\n\n".join(str(item.get("content") or "") for item in chunks if item.get("content"))
    return payload


def knowledge_items_for_chat(section_id: str | None = None) -> list[dict[str, Any]]:
    store = load_store()
    datasets = {str(item.get("id")): item for item in store.get("datasets", [])}
    items: list[dict[str, Any]] = []
    for record in store.get("records", []):
        dataset = datasets.get(str(record.get("dataset_id")))
        if not dataset:
            continue
        current_section = str(dataset.get("domainId") or "stories")
        if section_id and current_section != section_id:
            continue
        title = str(_record_field(record, "title") or dataset.get("name") or dataset.get("file_name") or "上传资料")
        content = str(_record_field(record, "content") or "")
        items.append({
            "id": f"platform-record-{record.get('id')}",
            "sectionId": current_section,
            "resourceType": "OCR识别文档" if record.get("text_kind") in {"ocr", "pdf-ocr"} else "上传文档文本" if dataset.get("dataset_kind") == "document" else "上传表格记录",
            "canonicalTitle": title,
            "translatedTitle": title,
            "year": _record_field(record, "publish_year") or 0,
            "language": "",
            "country": _record_field(record, "country") or "",
            "city": _record_field(record, "city") or "",
            "publisher": _record_field(record, "publisher") or "",
            "translator": _record_field(record, "translator") or "",
            "author": _record_field(record, "author") or "",
            "summary": content[:1200],
            "source": dataset.get("file_name") or dataset.get("name") or "管理员上传资料",
            "sourceKind": "uploaded-document" if dataset.get("dataset_kind") == "document" else "uploaded-table",
            "bookName": dataset.get("name") or "",
            "collectionId": f"dataset-{dataset.get('id')}",
            "relations": [
                f"属于数据集 -> {dataset.get('name') or dataset.get('file_name')}",
                f"归入子模块 -> {dataset.get('subModuleId')}",
            ],
            "searchText": f"{dataset.get('name') or dataset.get('file_name') or ''}\n{title}\n{content}",
            "graphNodeIds": [f"dataset:{dataset.get('id')}", f"platform-record:{record.get('id')}"],
        })
    return items


def _records_for_dataset(dataset_id: int | str) -> list[dict[str, Any]]:
    store = load_store()
    return [record for record in store.get("records", []) if str(record.get("dataset_id")) == str(dataset_id)]


def _records_for_submodule(sub_module_id: int | str) -> list[dict[str, Any]]:
    datasets = list_platform_datasets(sub_module_id)
    dataset_ids = {str(dataset["id"]) for dataset in datasets}
    dataset_meta = {str(dataset["id"]): dataset for dataset in datasets}
    store = load_store()
    records = []
    for record in store.get("records", []):
        dataset = dataset_meta.get(str(record.get("dataset_id")))
        if not dataset:
            continue
        enriched = dict(record)
        enriched["dataset_kind"] = dataset.get("dataset_kind") or "table"
        enriched["dataset"] = {
            "id": dataset.get("id"),
            "name": dataset.get("name"),
            "file_name": dataset.get("file_name"),
            "dataset_kind": dataset.get("dataset_kind") or "table",
            "text_kind": dataset.get("text_kind"),
            "status": dataset.get("status"),
            "description": dataset.get("description"),
            "analysis": dataset.get("analysis") or {},
        }
        records.append(enriched)
    return records or _builtin_records_for_submodule(sub_module_id)


def _record_field(record: dict[str, Any], field: str) -> Any:
    if field in record:
        return record.get(field)
    system = record.get("system") or {}
    if field in system:
        return system.get(field)
    raw = record.get("raw") or {}
    return raw.get(field)


def apply_filters(records: list[dict[str, Any]], filters: Any = None) -> list[dict[str, Any]]:
    if not filters:
        return records
    if isinstance(filters, str):
        try:
            filters = json.loads(filters)
        except Exception:
            return records
    if isinstance(filters, dict):
        filters = filters.get("conditions") or filters.get("filters") or [filters]
    if not isinstance(filters, list):
        return records

    def matches(record: dict[str, Any]) -> bool:
        for condition in filters:
            if not isinstance(condition, dict):
                continue
            field = str(condition.get("field") or condition.get("key") or "")
            op = str(condition.get("op") or condition.get("operator") or "eq").lower()
            expected = condition.get("value")
            actual = _record_field(record, field)
            actual_s = str(actual or "")
            expected_s = str(expected or "")
            if op in {"in", "one_of"}:
                if isinstance(expected, (list, tuple, set)):
                    expected_values = {str(item) for item in expected}
                else:
                    expected_values = {item.strip() for item in expected_s.replace("，", ",").split(",") if item.strip()}
                if actual_s not in expected_values:
                    return False
                continue
            if op in {"not_in", "not_one_of"}:
                if isinstance(expected, (list, tuple, set)):
                    expected_values = {str(item) for item in expected}
                else:
                    expected_values = {item.strip() for item in expected_s.replace("，", ",").split(",") if item.strip()}
                if actual_s in expected_values:
                    return False
                continue
            try:
                actual_n = float(actual_s)
                expected_n = float(expected_s)
            except Exception:
                actual_n = expected_n = None
            if op in {"eq", "=", "equals"} and actual_s != expected_s:
                return False
            if op in {"ne", "!=", "not_equals"} and actual_s == expected_s:
                return False
            if op in {"contains", "like"} and expected_s.lower() not in actual_s.lower():
                return False
            if op in {"not_contains", "not_like"} and expected_s.lower() in actual_s.lower():
                return False
            if op in {"gt", ">"} and not (actual_n is not None and actual_n > expected_n):
                return False
            if op in {"lt", "<"} and not (actual_n is not None and actual_n < expected_n):
                return False
            if op in {"gte", ">="} and not (actual_n is not None and actual_n >= expected_n):
                return False
            if op in {"lte", "<="} and not (actual_n is not None and actual_n <= expected_n):
                return False
            if op in {"between", "range"}:
                if isinstance(expected, (list, tuple)) and len(expected) >= 2:
                    start, end = expected[0], expected[1]
                else:
                    parts = [part for part in re.split(r"[,，~至-]", expected_s) if part.strip()]
                    start, end = (parts + [None, None])[:2]
                try:
                    start_n = float(start)
                    end_n = float(end)
                except Exception:
                    return False
                if not (actual_n is not None and min(start_n, end_n) <= actual_n <= max(start_n, end_n)):
                    return False
        return True

    return [record for record in records if matches(record)]


def paginate_records(
    records: list[dict[str, Any]],
    page: int = 1,
    page_size: int = 20,
    sort_by: str = "",
    sort_order: str = "asc",
    filters: Any = None,
) -> dict[str, Any]:
    filtered = apply_filters(records, filters)
    if sort_by:
        reverse = str(sort_order).lower() == "desc"
        filtered.sort(key=lambda record: str(_record_field(record, sort_by) or ""), reverse=reverse)
    page = max(1, int(page or 1))
    page_size = min(10000, max(1, int(page_size or 20)))
    start = (page - 1) * page_size
    rows = filtered[start:start + page_size]
    return {"records": rows, "total": len(filtered), "page": page, "page_size": page_size}


def dataset_records(dataset_id: int | str, **kwargs: Any) -> dict[str, Any]:
    return {**paginate_records(_records_for_dataset(dataset_id), **kwargs), "dataset": get_dataset(dataset_id)}


def submodule_records(sub_module_id: int | str, **kwargs: Any) -> dict[str, Any]:
    datasets = list_platform_datasets(sub_module_id)
    records = _records_for_submodule(sub_module_id)
    return {**paginate_records(records, **kwargs), "datasets": datasets or _builtin_dataset_metas(sub_module_id)}


def metrics_for_submodule(sub_module_id: int | str, filters: Any = None) -> dict[str, Any]:
    records = apply_filters(_records_for_submodule(sub_module_id), filters)
    translators = {(_record_field(record, "translator") or "").strip() for record in records if (_record_field(record, "translator") or "").strip()}
    publishers = {(_record_field(record, "publisher") or "").strip() for record in records if (_record_field(record, "publisher") or "").strip()}
    countries = {(_record_field(record, "country") or "").strip() for record in records if (_record_field(record, "country") or "").strip()}
    return {
        "total_documents": len(records),
        "total_translators": len(translators),
        "total_publishers": len(publishers),
        "total_countries": len(countries),
        "trend": [max(0, len(records) - 3), max(0, len(records) - 1), len(records)],
    }


def knowledge_graph(sub_module_id: int | str, filters: Any = None) -> dict[str, Any]:
    records = apply_filters(_records_for_submodule(sub_module_id), filters)
    nodes: dict[str, dict[str, Any]] = {}
    edge_counter: Counter[tuple[str, str, str]] = Counter()
    type_meta = {
        "author": ("作者", "#165DFF", 24),
        "translator": ("译者", "#00B42A", 20),
        "title": ("作品", "#FF7D00", 16),
        "publisher": ("出版社", "#F53F3F", 12),
        "city": ("出版地", "#722ED1", 12),
    }
    for record in records:
        ids = {}
        for field, (label, color, size) in type_meta.items():
            value = str(_record_field(record, field) or "").strip()
            if not value:
                continue
            node_id = f"{field}:{value}"
            ids[field] = node_id
            nodes.setdefault(node_id, {"id": node_id, "name": value, "label": value, "type": label, "symbolSize": size, "itemStyle": {"color": color}, "count": 0})
            nodes[node_id]["count"] += 1
        title_id = ids.get("title")
        for field in ["author", "translator", "publisher", "city"]:
            if title_id and ids.get(field):
                edge_counter[(ids[field], title_id, "关联")] += 1
    edges = [
        {"source": source, "target": target, "label": relation, "value": count, "lineStyle": {"width": min(4, max(1, count))}}
        for (source, target, relation), count in edge_counter.items()
    ]
    return {"nodes": list(nodes.values()), "edges": edges, "links": edges}


def map_data(sub_module_id: int | str, map_type: str = "publication", filters: Any = None) -> dict[str, Any]:
    records = apply_filters(_records_for_submodule(sub_module_id), filters)
    counts: Counter[str] = Counter()
    cities: Counter[str] = Counter()
    routes: Counter[tuple[str, str]] = Counter()
    for record in records:
        country = str(_record_field(record, "country") or "").strip()
        city = str(_record_field(record, "city") or "").strip()
        source = str(_record_field(record, "source") or "").strip() or "China"
        if country:
            counts[country] += 1
        if city:
            cities[city] += 1
        if map_type == "route" and source and country:
            routes[(source, country)] += 1
    return {
        "map_type": map_type,
        "countries": [{"name": name, "value": value} for name, value in counts.most_common()],
        "cities": [{"name": name, "value": value} for name, value in cities.most_common()],
        "routes": [{"from": source, "to": target, "value": value} for (source, target), value in routes.most_common()],
    }


def _tokenize(text: str) -> list[str]:
    try:
        import jieba  # type: ignore

        words = list(jieba.cut(text))
    except Exception:
        words = re.findall(r"[\u4e00-\u9fff]{2,}|[A-Za-z][A-Za-z\-']{2,}", text)
    result = []
    for word in words:
        token = str(word).strip().lower()
        if len(token) < 2 or token in STOPWORDS:
            continue
        result.append(token)
    return result


DOCUMENT_SCRIPT_META = [
    ("zh", "汉字", r"[\u4e00-\u9fff]"),
    ("latin", "拉丁", r"[A-Za-zÀ-ž]"),
    ("cyrillic", "西里尔", r"[\u0400-\u04ff]"),
    ("greek", "希腊", r"[\u0370-\u03ff]"),
    ("kana", "假名", r"[\u3040-\u30ff]"),
    ("hangul", "韩文", r"[\uac00-\ud7af]"),
    ("arabic", "阿拉伯", r"[\u0600-\u06ff]"),
]

DOCUMENT_TOPIC_RULES = [
    {"id": "concept", "label": "核心概念", "seeds": ["文学", "文化", "叙事", "传统", "典籍", "经典", "思想", "文本", "文献", "知识"]},
    {"id": "translation", "label": "译介传播", "seeds": ["翻译", "译者", "译本", "传播", "出版", "接受", "改写", "编译", "介绍", "海外"]},
    {"id": "method", "label": "方法证据", "seeds": ["研究", "分析", "比较", "统计", "方法", "材料", "版本", "来源", "注释", "考证"]},
    {"id": "actor", "label": "主体机构", "seeds": ["作者", "读者", "学者", "出版社", "大学", "机构", "国家", "社会", "民族", "编者"]},
    {"id": "form", "label": "体裁形式", "seeds": ["小说", "诗歌", "戏剧", "故事", "序跋", "论文", "章节", "卷", "篇", "集"]},
]


def _script_counts(text: str) -> tuple[list[dict[str, Any]], str]:
    value = str(text or "")
    items = []
    used = 0
    for key, label, pattern in DOCUMENT_SCRIPT_META:
        count = len(re.findall(pattern, value))
        used += count
        items.append({"key": key, "label": label, "count": count})
    total = max(1, used)
    for item in items:
        item["ratio"] = round(item["count"] / total, 4)
    dominant = max(items, key=lambda item: item["count"], default={"key": "all", "label": "未识别"})
    return items, str(dominant.get("key") or "all")


def _token_script(token: str) -> str:
    value = str(token or "")
    for key, _label, pattern in DOCUMENT_SCRIPT_META:
        if re.search(pattern, value):
            return key
    return "other"


def _script_label(script: str) -> str:
    return next((label for key, label, _pattern in DOCUMENT_SCRIPT_META if key == script), "其他")


def _document_topic_for(term: str) -> str:
    value = str(term or "").lower()
    for rule in DOCUMENT_TOPIC_RULES:
        for seed in rule["seeds"]:
            seed_value = str(seed).lower()
            if seed_value and (seed_value == value or seed_value in value or value in seed_value):
                return str(rule["id"])
    script = _token_script(value)
    if script == "zh":
        return "concept"
    if script in {"latin", "greek", "cyrillic"}:
        return "translation"
    return "method"


def _document_topic_label(topic: str) -> str:
    return next((str(rule["label"]) for rule in DOCUMENT_TOPIC_RULES if rule["id"] == topic), "综合术语")


def _normalize_document_token(token: Any) -> str:
    value = re.sub(r"\s+", "", str(token or "")).strip(".,!?;:，。！？；：《》“”‘’（）()[]【】\"'")
    return value.lower() if re.search(r"[A-Za-zÀ-ž]", value) and not re.search(r"[\u4e00-\u9fff]", value) else value


def _valid_document_token(token: str, script_scope: str = "all") -> bool:
    value = _normalize_document_token(token)
    if len(value) < 2 or len(value) > 32:
        return False
    if value.lower() in STOPWORDS or value in STOPWORDS:
        return False
    if re.fullmatch(r"\d+|[\W_]+", value, flags=re.U):
        return False
    script = _token_script(value)
    if script_scope and script_scope != "all" and script != script_scope:
        return False
    if script == "zh":
        if len(value) < 2:
            return False
        if re.fullmatch(r"[一二三四五六七八九十百千万年月日章节卷部篇]+", value):
            return False
    elif script == "latin" and (len(value) < 4 or re.fullmatch(r"[ivxlcdm]+", value)):
        return False
    return bool(re.search(r"[\u4e00-\u9fffA-Za-zÀ-ž\u0400-\u04ff\u0370-\u03ff\u3040-\u30ff\uac00-\ud7af\u0600-\u06ff]", value))


def _document_tokens(text: str, script_scope: str = "all") -> list[str]:
    value = unicodedata.normalize("NFKC", str(text or ""))
    tokens: list[str] = []
    try:
        import jieba.posseg as pseg  # type: ignore

        for word, flag in pseg.cut(value):
            clean = _normalize_document_token(word)
            pos = str(flag or "").lower()
            if not _valid_document_token(clean, script_scope):
                continue
            if _token_script(clean) == "zh" and not (pos.startswith(("n", "v", "a", "eng")) or clean in {"文学", "文化", "翻译", "传播", "叙事", "经典", "典籍"}):
                continue
            tokens.append(clean)
    except Exception:
        pattern = r"[\u4e00-\u9fff]{2,8}|[A-Za-zÀ-ž][A-Za-zÀ-ž'\-]{2,}|[\u0400-\u04ff]{2,}|[\u0370-\u03ff]{2,}|[\u3040-\u30ff]{2,}|[\uac00-\ud7af]{2,}|[\u0600-\u06ff]{2,}"
        tokens = [_normalize_document_token(token) for token in re.findall(pattern, value)]
    return [token for token in tokens if _valid_document_token(token, script_scope)]


def _document_textrank(token_docs: list[list[str]], window: int = 4, iterations: int = 24) -> Counter[str]:
    graph: dict[str, Counter[str]] = defaultdict(Counter)
    for tokens in token_docs:
        sequence = [token for token in tokens if _valid_document_token(token)]
        for index, source in enumerate(sequence):
            for distance, target in enumerate(sequence[index + 1:index + window], start=1):
                if source == target:
                    continue
                weight = 1.0 / distance
                graph[source][target] += weight
                graph[target][source] += weight
    terms = set(graph)
    for neighbors in graph.values():
        terms.update(neighbors)
    if not terms:
        return Counter()
    scores = {term: 1.0 for term in terms}
    damping = 0.85
    for _ in range(iterations):
        next_scores = {term: 1.0 - damping for term in terms}
        for source, neighbors in graph.items():
            total_weight = sum(neighbors.values()) or 1.0
            base = scores.get(source, 1.0) / total_weight
            for target, weight in neighbors.items():
                next_scores[target] += damping * base * weight
        scores = next_scores
    return Counter(scores)


def _normalize_scores(scores: Counter[str] | dict[str, float]) -> dict[str, float]:
    if not scores:
        return {}
    max_score = max(float(value) for value in scores.values()) or 1.0
    return {term: float(value) / max_score for term, value in scores.items()}


def _analysis_text_for_record(record: dict[str, Any]) -> str:
    return clean_text(" ".join(
        str(_record_field(record, field) or "")
        for field in ["content", "preface", "theme", "title", "notes"]
    ))


def _dataset_id_for_record(record: dict[str, Any]) -> str:
    return str(record.get("dataset_id") or record.get("dataset", {}).get("id") or record.get("id") or "")


def document_text_analysis(
    sub_module_id: int | str,
    scope: str = "module",
    language_scope: str = "all",
    document_ids: list[str] | None = None,
    top_n: int = 120,
    filters: Any = None,
) -> dict[str, Any]:
    records = apply_filters(_records_for_submodule(sub_module_id), filters)
    wanted_ids = {str(item) for item in (document_ids or []) if str(item).strip()}
    if wanted_ids:
        records = [
            record for record in records
            if str(record.get("id")) in wanted_ids
            or _dataset_id_for_record(record) in wanted_ids
            or str(record.get("dataset", {}).get("id") or "") in wanted_ids
        ]
    dataset_titles: dict[str, str] = {}
    doc_texts: dict[str, list[str]] = defaultdict(list)
    doc_records: dict[str, list[dict[str, Any]]] = defaultdict(list)
    for record in records:
        dataset = record.get("dataset") or {}
        dataset_id = _dataset_id_for_record(record)
        if not dataset_id:
            continue
        dataset_titles.setdefault(dataset_id, str(dataset.get("name") or dataset.get("file_name") or _record_field(record, "title") or f"文档 {dataset_id}"))
        text = _analysis_text_for_record(record)
        if text.strip():
            doc_texts[dataset_id].append(text)
            doc_records[dataset_id].append(record)

    if scope == "document" and wanted_ids:
        doc_texts = defaultdict(list, {key: value for key, value in doc_texts.items() if key in wanted_ids})
    docs = [
        {
            "id": dataset_id,
            "title": dataset_titles.get(dataset_id, f"文档 {index + 1}"),
            "text": "\n\n".join(parts),
            "records": doc_records.get(dataset_id, []),
        }
        for index, (dataset_id, parts) in enumerate(doc_texts.items())
        if "\n".join(parts).strip()
    ]
    text_scope = language_scope if language_scope in {key for key, _label, _pattern in DOCUMENT_SCRIPT_META} else "all"
    token_docs: list[list[str]] = []
    total_counts: Counter[str] = Counter()
    doc_frequency: Counter[str] = Counter()
    script_counter: Counter[str] = Counter()
    document_metrics: list[dict[str, Any]] = []
    chunk_windows: list[list[str]] = []

    for doc in docs:
        tokens = _document_tokens(doc["text"], text_scope)
        token_docs.append(tokens)
        counts = Counter(tokens)
        total_counts.update(counts)
        doc_frequency.update(counts.keys())
        profile, dominant = _script_counts(doc["text"])
        for item in profile:
            script_counter[item["key"]] += int(item["count"])
        paragraphs = [part for part in re.split(r"\n{2,}|\r?\n", doc["text"]) if part.strip()]
        document_metrics.append({
            "id": doc["id"],
            "title": doc["title"],
            "charCount": len(doc["text"]),
            "paragraphs": len(paragraphs),
            "tokenCount": len(tokens),
            "uniqueTerms": len(counts),
            "lexicalDiversity": round(len(counts) / max(1, len(tokens)), 4),
            "dominantScript": dominant,
            "scriptProfile": profile,
        })
        chunks = [
            str(_record_field(record, "content") or "")
            for record in doc.get("records") or []
            if str(_record_field(record, "content") or "").strip()
        ] or paragraphs or [doc["text"]]
        for chunk in chunks[:160]:
            chunk_tokens = _document_tokens(chunk, text_scope)
            if chunk_tokens:
                chunk_windows.append(chunk_tokens)

    doc_count = max(1, len(token_docs))
    tfidf_scores: Counter[str] = Counter()
    for tokens in token_docs:
        counts = Counter(tokens)
        if not counts:
            continue
        max_count = max(counts.values()) or 1
        for term, count in counts.items():
            idf = math.log((1 + doc_count) / (1 + doc_frequency[term])) + 1.0
            tf = math.log1p(count) / math.log1p(max_count)
            tfidf_scores[term] += tf * idf
    textrank_scores = _document_textrank(token_docs)
    tfidf_norm = _normalize_scores(tfidf_scores)
    textrank_norm = _normalize_scores(textrank_scores)
    max_count = max(total_counts.values(), default=1)
    fused_scores: dict[str, float] = {}
    for term in set(total_counts) | set(tfidf_scores) | set(textrank_scores):
        script = _token_script(term)
        length_bonus = min(1.22, 0.86 + min(len(term), 8) / 24)
        topic_bonus = 1.12 if _document_topic_for(term) in {"concept", "translation"} else 1.0
        frequency = math.log1p(total_counts.get(term, 0)) / math.log1p(max_count)
        fused_scores[term] = (
            0.44 * tfidf_norm.get(term, 0.0)
            + 0.30 * textrank_norm.get(term, 0.0)
            + 0.18 * frequency
            + 0.08 * (1.0 if script == "zh" else 0.92)
        ) * length_bonus * topic_bonus
    ranked_terms = sorted(
        (term for term in fused_scores if _valid_document_token(term, text_scope)),
        key=lambda term: (fused_scores[term], total_counts.get(term, 0), len(term)),
        reverse=True,
    )
    selected: list[str] = []
    for term in ranked_terms:
        if any(term == existing for existing in selected):
            continue
        if any(term in existing and len(term) <= 3 and fused_scores[term] <= fused_scores[existing] * 1.08 for existing in selected):
            continue
        selected = [
            existing for existing in selected
            if not (existing in term and len(existing) <= 3 and fused_scores[existing] <= fused_scores[term] * 1.08)
        ]
        selected.append(term)
        if len(selected) >= max(20, min(600, int(top_n or 120))):
            break
    max_score = max((fused_scores.get(term, 0.0) for term in selected), default=1.0) or 1.0
    keywords = [
        {
            "word": term,
            "text": term,
            "name": term,
            "count": int(total_counts.get(term, 0)),
            "docCount": int(doc_frequency.get(term, 0)),
            "score": round(fused_scores.get(term, 0.0) * 100, 4),
            "value": round((fused_scores.get(term, 0.0) / max_score) * 100, 3),
            "tfidf": round(tfidf_norm.get(term, 0.0), 6),
            "textrank": round(textrank_norm.get(term, 0.0), 6),
            "script": _token_script(term),
            "language": _script_label(_token_script(term)),
            "topic": _document_topic_for(term),
            "topicLabel": _document_topic_label(_document_topic_for(term)),
        }
        for term in selected
    ]

    important = {item["word"] for item in keywords[:36]}
    cooccurrence: Counter[tuple[str, str]] = Counter()
    for tokens in chunk_windows:
        window = [token for token in tokens if token in important]
        for index, left in enumerate(window):
            for right in window[index + 1:index + 8]:
                if left != right:
                    cooccurrence[tuple(sorted((left, right)))] += 1
    scripts_total = max(1, sum(script_counter.values()))
    script_distribution = [
        {
            "key": key,
            "label": label,
            "count": int(script_counter.get(key, 0)),
            "ratio": round(script_counter.get(key, 0) / scripts_total, 4),
        }
        for key, label, _pattern in DOCUMENT_SCRIPT_META
    ]
    return {
        "subModuleId": str(sub_module_id),
        "scope": scope or "module",
        "languageScope": language_scope or "all",
        "documentIds": list(wanted_ids),
        "documentCount": len(docs),
        "recordCount": len(records),
        "tokenCount": sum(len(tokens) for tokens in token_docs),
        "uniqueTermCount": len(total_counts),
        "keywords": keywords[:max(20, min(240, int(top_n or 120)))],
        "wordCloud": keywords[:max(60, min(600, int(top_n or 120) * 4))],
        "cooccurrence": [
            {"source": source, "target": target, "count": int(count), "value": int(count)}
            for (source, target), count in cooccurrence.most_common(80)
        ],
        "scriptDistribution": script_distribution,
        "languageOptions": [{"value": "all", "label": "全文所有语种"}] + [
            {"value": key, "label": label, "count": int(script_counter.get(key, 0))}
            for key, label, _pattern in DOCUMENT_SCRIPT_META
            if script_counter.get(key, 0) > 0
        ],
        "documents": document_metrics,
        "topics": [
            {
                "id": rule["id"],
                "label": rule["label"],
                "value": sum(item["value"] for item in keywords[:160] if item["topic"] == rule["id"]),
                "count": sum(item["count"] for item in keywords[:160] if item["topic"] == rule["id"]),
            }
            for rule in DOCUMENT_TOPIC_RULES
        ],
        "method": "multilingual-tfidf-textrank-wordcloud",
    }


def _advanced_text_documents_for_records(records: list[dict[str, Any]]) -> list[dict[str, Any]]:
    grouped: dict[str, dict[str, Any]] = {}
    loose_index = 0
    for record in records:
        dataset = record.get("dataset") or {}
        dataset_id = _dataset_id_for_record(record)
        dataset_kind = str(record.get("dataset_kind") or dataset.get("dataset_kind") or "").lower()
        is_document_dataset = dataset_kind == "document" or bool(dataset.get("text_kind"))
        if is_document_dataset and dataset_id:
            item = grouped.setdefault(dataset_id, {
                "id": dataset_id,
                "title": str(dataset.get("name") or dataset.get("file_name") or _record_field(record, "title") or f"文档 {dataset_id}"),
                "filename": dataset.get("file_name") or "",
                "source": dataset.get("file_name") or "",
                "author": _record_field(record, "author") or _record_field(record, "translator") or "",
                "translator": _record_field(record, "translator") or "",
                "year": _record_field(record, "publish_year") or "",
                "language": _record_field(record, "language") or "",
                "textParts": [],
                "metadata": {
                    "datasetId": dataset_id,
                    "datasetKind": dataset_kind or "document",
                    "recordIds": [],
                },
            })
            text = _analysis_text_for_record(record)
            if text.strip():
                item["textParts"].append(text)
            item["metadata"]["recordIds"].append(record.get("id"))
            continue

        loose_index += 1
        text = _analysis_text_for_record(record)
        if not text.strip():
            continue
        record_id = str(record.get("id") or f"record-{loose_index}")
        grouped[f"record:{record_id}"] = {
            "id": record_id,
            "title": str(_record_field(record, "title") or f"记录 {loose_index}"),
            "filename": str(_record_field(record, "source") or ""),
            "source": str(_record_field(record, "source") or ""),
            "author": _record_field(record, "author") or _record_field(record, "translator") or "",
            "translator": _record_field(record, "translator") or "",
            "year": _record_field(record, "publish_year") or "",
            "language": _record_field(record, "language") or "",
            "textParts": [text],
            "metadata": {
                "recordId": record_id,
                "datasetId": dataset_id,
                "datasetKind": dataset_kind or "table",
            },
        }
    documents = []
    for item in grouped.values():
        text = "\n\n".join(part for part in item.pop("textParts", []) if str(part).strip()).strip()
        if not text:
            continue
        item["text"] = text
        documents.append(item)
    return documents


def advanced_text_visuals(
    sub_module_id: int | str,
    scope: str = "single",
    document_id: str = "",
    query: str = "",
    method_id: str = "semantic-manifold",
    topic_count: int = 18,
    filters: Any = None,
) -> dict[str, Any]:
    records = apply_filters(_records_for_submodule(sub_module_id), filters)
    documents = _advanced_text_documents_for_records(records)
    submodule = find_submodule(sub_module_id) or {}
    title = str(submodule.get("name") or f"子模块 {sub_module_id}")
    return advanced_text_visualization_payload(
        documents,
        source_id=f"submodule-{sub_module_id}",
        corpus_title=title,
        document_id=document_id,
        scope=scope,
        query=query,
        method_id=method_id,
        topic_count=topic_count,
    )


def word_frequency(sub_module_id: int | str, text_fields: list[str] | None = None, top_n: int = 20, filters: Any = None) -> dict[str, Any]:
    records = apply_filters(_records_for_submodule(sub_module_id), filters)
    fields = text_fields or ["content", "preface", "theme", "title"]
    counter: Counter[str] = Counter()
    for record in records:
        text = " ".join(str(_record_field(record, field) or "") for field in fields)
        counter.update(_tokenize(text))
    return {"items": [{"word": word, "count": count, "name": word, "value": count} for word, count in counter.most_common(top_n)]}


def time_evolution(sub_module_id: int | str, time_field: str = "publish_year", aggregation: str = "year", filters: Any = None) -> dict[str, Any]:
    records = apply_filters(_records_for_submodule(sub_module_id), filters)
    counter: Counter[str] = Counter()
    for record in records:
        value = str(_record_field(record, time_field) or _record_field(record, "publish_year") or "")
        match = re.search(r"\d{3,4}", value)
        if not match:
            continue
        year = int(match.group())
        if aggregation == "decade":
            key = f"{year // 10 * 10}s"
        elif aggregation == "century":
            key = f"{(year - 1) // 100 + 1}世纪"
        else:
            key = str(year)
        counter[key] += 1
    return {"series": [{"time": key, "value": counter[key]} for key in sorted(counter.keys())]}


def topic_clustering(sub_module_id: int | str, text_field: str = "content", n_topics: int = 5, filters: Any = None) -> dict[str, Any]:
    records = apply_filters(_records_for_submodule(sub_module_id), filters)
    docs = [str(_record_field(record, text_field) or _record_field(record, "content") or "") for record in records]
    docs = [doc for doc in docs if doc.strip()]
    if not docs:
        return {"topics": [], "documents": []}
    token_docs = [Counter(_tokenize(doc)) for doc in docs]
    token_docs = [counts for counts in token_docs if counts]
    if not token_docs:
        return {"topics": [], "documents": []}

    global_counts: Counter[str] = Counter()
    document_frequency: Counter[str] = Counter()
    for counts in token_docs:
        global_counts.update(counts)
        document_frequency.update(counts.keys())

    topic_count = min(max(1, int(n_topics or 1)), max(1, len(global_counts)))
    seeds = [
        word
        for word, _count in global_counts.most_common(max(20, topic_count * 6))
        if document_frequency[word] >= 1
    ][:topic_count]
    if not seeds:
        return {"topics": [], "documents": []}

    distribution: list[list[float]] = []
    topic_doc_counts = [0 for _ in seeds]
    topic_terms: list[Counter[str]] = [Counter() for _ in seeds]
    for doc_index, counts in enumerate(token_docs):
        scores = []
        for seed_index, seed in enumerate(seeds):
            rank_bonus = max(0.0, (len(seeds) - seed_index) / max(1, len(seeds)) * 0.01)
            scores.append(float(counts.get(seed, 0)) + rank_bonus)
        best = max(range(len(scores)), key=lambda index: scores[index]) if scores else doc_index % len(seeds)
        if scores and scores[best] <= 0.01:
            best = doc_index % len(seeds)
        topic_doc_counts[best] += 1
        topic_terms[best].update(counts)
        total = sum(scores) or 1.0
        row = [score / total for score in scores]
        if not any(row):
            row[best] = 1.0
        distribution.append(row)

    topics = []
    for index, seed in enumerate(seeds):
        keywords = [word for word, _count in topic_terms[index].most_common(5)]
        if seed not in keywords:
            keywords.insert(0, seed)
        topics.append({
            "id": index,
            "keywords": keywords[:5] or [seed],
            "size": max(1, topic_doc_counts[index]),
            "documentCount": topic_doc_counts[index],
        })
    return {"topics": topics, "documents": distribution, "method": "frequency-seed-clustering"}


def comparison(sub_module_id: int | str, dimensions: list[str] | None = None, filters: Any = None) -> dict[str, Any]:
    records = apply_filters(_records_for_submodule(sub_module_id), filters)
    dims = dimensions or ["country", "publisher", "translator"]
    result = []
    for dim in dims:
        counter = Counter(str(_record_field(record, dim) or "未注明").strip() or "未注明" for record in records)
        result.append({"dimension": dim, "items": [{"name": name, "value": value} for name, value in counter.most_common(20)]})
    return {"dimensions": result}


def word_distance(sub_module_id: int | str, word_a: str, word_b: str, max_distance: int = 20, text_field: str = "content", filters: Any = None) -> dict[str, Any]:
    records = apply_filters(_records_for_submodule(sub_module_id), filters)
    results = []
    pattern = re.compile(f"({re.escape(word_a)})(.{{0,{max_distance}}})({re.escape(word_b)})|({re.escape(word_b)})(.{{0,{max_distance}}})({re.escape(word_a)})", re.I | re.S)
    for record in records:
        text = str(_record_field(record, text_field) or _record_field(record, "content") or "")
        for match in pattern.finditer(text):
            start = max(0, match.start() - 30)
            end = min(len(text), match.end() + 30)
            results.append({
                "record_id": record.get("id"),
                "title": _record_field(record, "title"),
                "snippet": text[start:end],
                "distance": len(match.group(2) or match.group(5) or ""),
            })
            if len(results) >= 100:
                return {"items": results}
    return {"items": results}


def word_trend(sub_module_id: int | str, words: list[str], time_field: str = "publish_year", filters: Any = None) -> dict[str, Any]:
    records = apply_filters(_records_for_submodule(sub_module_id), filters)
    series: dict[str, Counter[str]] = {word: Counter() for word in words}
    for record in records:
        value = str(_record_field(record, time_field) or _record_field(record, "publish_year") or "")
        match = re.search(r"\d{3,4}", value)
        if not match:
            continue
        year = match.group()
        text = str(_record_field(record, "content") or "") + " " + str(_record_field(record, "title") or "")
        for word in words:
            count = text.count(word)
            if count:
                series[word][year] += count
    return {"series": [{"word": word, "data": [{"time": year, "value": count} for year, count in sorted(counter.items())]} for word, counter in series.items()]}


def full_text_search(keyword: str, filters: Any = None) -> dict[str, Any]:
    store = load_store()
    records = apply_filters(store.get("records", []), filters)
    keyword_lower = keyword.lower()
    matches = []
    for record in records:
        text = " ".join(str(_record_field(record, field) or "") for field in ["title", "author", "translator", "publisher", "country", "city", "theme", "content", "preface", "notes"])
        if keyword_lower not in text.lower():
            continue
        index = text.lower().find(keyword_lower)
        snippet = text[max(0, index - 50): index + len(keyword) + 80]
        matches.append({"record_id": record.get("id"), "dataset_id": record.get("dataset_id"), "title": _record_field(record, "title"), "snippet": snippet.replace(keyword, f"<mark>{keyword}</mark>")})
    return {"items": matches, "total": len(matches)}


def admin_submodules() -> dict[str, Any]:
    return platform_registry()


def create_submodule(payload: dict[str, Any]) -> dict[str, Any]:
    store = load_store()
    domain = find_domain(payload.get("knowledge_domain_id") or payload.get("domainId") or payload.get("domain_id") or "classics")
    if not domain:
        raise ValueError("knowledge_domain_id is invalid.")
    existing_ids = {str(item.get("id")) for item in all_submodules()}
    timestamp_ms = int(time.time() * 1000)
    submodule_id = f"{domain['id']}-{timestamp_ms}"
    suffix = 1
    while submodule_id in existing_ids:
        suffix += 1
        submodule_id = f"{domain['id']}-{timestamp_ms}-{suffix}"
    submodule = {
        "id": submodule_id,
        "numericId": max([item["numericId"] for item in all_submodules()] + [1000]) + int(time.time()) % 100000,
        "knowledge_domain_id": domain["numericId"],
        "domainId": domain["id"],
        "name": str(payload.get("name") or "").strip(),
        "description": str(payload.get("description") or ""),
        "type": str(payload.get("type") or "topic"),
        "language": payload.get("language"),
        "enabled_components": payload.get("enabled_components") or payload.get("enabledComponents") or [],
        "sort_order": int(payload.get("sort_order") or 999),
        "is_active": True,
        "dynamic": True,
    }
    if not submodule["name"]:
        raise ValueError("Submodule name is required.")
    store["deleted_submodule_ids"] = [item for item in store.get("deleted_submodule_ids", []) if str(item) != str(submodule["id"])]
    for item in store["domains"]:
        if item["id"] == domain["id"]:
            item.setdefault("submodules", []).append(submodule)
            break
    save_store(store)
    record_operation("submodule_create", f"Created submodule {submodule['name']}")
    return submodule


def update_submodule(submodule_id: str, payload: dict[str, Any]) -> dict[str, Any]:
    store = load_store()
    for domain in store["domains"]:
        for submodule in domain.get("submodules", []):
            if str(submodule.get("id")) == str(submodule_id) or str(submodule.get("numericId")) == str(submodule_id):
                for key in ["name", "description", "type", "language", "sort_order", "is_active"]:
                    if key in payload:
                        submodule[key] = payload[key]
                if "enabled_components" in payload:
                    submodule["enabled_components"] = payload["enabled_components"]
                if "enabledComponents" in payload:
                    submodule["enabled_components"] = payload["enabledComponents"]
                save_store(store)
                record_operation("submodule_update", f"Updated submodule {submodule_id}")
                return submodule
    raise KeyError("Submodule not found.")


def delete_submodule(submodule_id: str) -> None:
    store = load_store()
    for domain in store["domains"]:
        before = len(domain.get("submodules", []))
        removed = [
            item for item in domain.get("submodules", [])
            if str(item.get("id")) == str(submodule_id) or str(item.get("numericId")) == str(submodule_id)
        ]
        domain["submodules"] = [item for item in domain.get("submodules", []) if str(item.get("id")) != str(submodule_id) and str(item.get("numericId")) != str(submodule_id)]
        if len(domain["submodules"]) != before:
            deleted = store.setdefault("deleted_submodule_ids", [])
            deleted.append(str(submodule_id))
            for item in removed:
                deleted.extend([str(item.get("id")), str(item.get("numericId"))])
            store["deleted_submodule_ids"] = sorted(set(store["deleted_submodule_ids"]))
            save_store(store)
            record_operation("submodule_delete", f"Deleted submodule {submodule_id}")
            return
    raise KeyError("Submodule not found.")


def update_dataset(dataset_id: str | int, payload: dict[str, Any]) -> dict[str, Any]:
    store = load_store()
    for dataset in store.get("datasets", []):
        if str(dataset.get("id")) == str(dataset_id):
            for key in ["name", "file_name", "status", "error_message", "force_ocr", "dataset_kind"]:
                if key in payload:
                    dataset[key] = payload[key]
            for key in ["description", "affected_pages", "key_fields"]:
                if key in payload:
                    dataset[key] = payload[key]
            dataset["updated_at"] = _now_iso()
            save_store(store)
            record_operation("dataset_update", f"Updated dataset {dataset_id}")
            return dataset
    raise KeyError("Dataset not found.")


def delete_platform_dataset(dataset_id: str | int) -> None:
    store = load_store()
    dataset = next((item for item in store.get("datasets", []) if str(item.get("id")) == str(dataset_id)), None)
    if not dataset:
        raise KeyError("Dataset not found.")
    store["datasets"] = [item for item in store["datasets"] if str(item.get("id")) != str(dataset_id)]
    store["records"] = [item for item in store.get("records", []) if str(item.get("dataset_id")) != str(dataset_id)]
    store["field_mappings"] = [item for item in store.get("field_mappings", []) if str(item.get("dataset_id")) != str(dataset_id)]
    save_store(store)
    path = Path(dataset.get("file_path") or "")
    if path.exists():
        try:
            path.unlink()
            if path.parent != UPLOAD_ROOT:
                path.parent.rmdir()
        except OSError:
            pass
    record_operation("dataset_delete", f"Deleted dataset {dataset_id}")


def reparse_dataset(dataset_id: str | int) -> dict[str, Any]:
    dataset = get_dataset(dataset_id)
    path = Path(dataset.get("file_path") or "")
    if not path.exists():
        raise FileNotFoundError("Original uploaded file is missing.")
    file_bytes = path.read_bytes()
    store = load_store()
    store["records"] = [item for item in store.get("records", []) if str(item.get("dataset_id")) != str(dataset_id)]
    save_store(store)
    mappings = {item["file_column"]: item["system_field"] for item in load_store().get("field_mappings", []) if str(item.get("dataset_id")) == str(dataset_id)}
    delete_platform_dataset(dataset_id)
    return upload_dataset_file(
        dataset["subModuleId"],
        file_bytes,
        dataset["file_name"],
        mappings,
        dataset.get("name") or Path(dataset.get("file_name") or "").stem,
        dataset.get("description") or "",
        dataset.get("affected_pages") or [],
        dataset.get("key_fields") or [],
        dataset.get("uploaded_by"),
        dataset.get("dataset_kind") or "",
        bool(dataset.get("force_ocr")),
    )


def clear_cache(sub_module_id: str | int | None = None) -> dict[str, Any]:
    store = load_store()
    before = len(store.get("visualization_cache", []))
    if sub_module_id in (None, "", "all"):
        store["visualization_cache"] = []
    else:
        store["visualization_cache"] = [item for item in store.get("visualization_cache", []) if str(item.get("sub_module_id")) != str(sub_module_id)]
    save_store(store)
    return {"ok": True, "removed": before - len(store.get("visualization_cache", []))}


def system_config() -> dict[str, Any]:
    return load_store().get("system_config", {})


def update_system_config(payload: dict[str, Any]) -> dict[str, Any]:
    store = load_store()
    current = store.setdefault("system_config", {})
    for key, value in payload.items():
        if isinstance(value, dict) and isinstance(current.get(key), dict):
            current[key].update(value)
        else:
            current[key] = value
    save_store(store)
    record_operation("system_config_update", "Updated system configuration")
    return current


def _prune_old_backups(root: Path, pattern: str) -> None:
    cutoff = datetime.now() - timedelta(days=RETENTION_DAYS)
    for path in root.glob(pattern):
        try:
            if datetime.fromtimestamp(path.stat().st_mtime) < cutoff:
                if path.is_dir():
                    shutil.rmtree(path, ignore_errors=True)
                else:
                    path.unlink(missing_ok=True)
        except Exception:
            continue


def _add_backup_job(kind: str, status: str, path: str, message: str = "") -> dict[str, Any]:
    store = load_store()
    item = {"id": store.get("next_backup_id", 1), "kind": kind, "status": status, "path": path, "message": message, "created_at": _now_iso()}
    store["next_backup_id"] = item["id"] + 1
    store.setdefault("backup_jobs", []).append(item)
    save_store(store)
    record_operation(f"backup_{kind}", f"{status}: {path} {message}".strip())
    return item


def database_backup(manual: bool = True) -> dict[str, Any]:
    ensure_dirs()
    target = DATABASE_BACKUP_ROOT / f"backup_{_timestamp()}.sql"
    database_url = os.environ.get("DATABASE_URL", "").strip()
    if database_url and shutil.which("pg_dump"):
        try:
            with target.open("wb") as handle:
                subprocess.run(["pg_dump", database_url], stdout=handle, stderr=subprocess.PIPE, check=True, timeout=3600)
            status = "completed"
            message = "pg_dump completed"
        except Exception as error:
            status = "failed"
            message = str(error)
            target.write_text(f"-- Backup failed: {message}\n", encoding="utf-8")
    else:
        snapshot = {"created_at": _now_iso(), "source": "json-compatible-store", "store": load_store()}
        target.write_text(json.dumps(snapshot, ensure_ascii=False, indent=2), encoding="utf-8")
        status = "completed"
        message = "DATABASE_URL or pg_dump unavailable; exported JSON-compatible SQL placeholder."
    _prune_old_backups(DATABASE_BACKUP_ROOT, "backup_*.sql")
    return _add_backup_job("database", status, str(target), message)


def files_backup(manual: bool = True) -> dict[str, Any]:
    ensure_dirs()
    target = FILE_BACKUP_ROOT / f"files_{_timestamp()}"
    if shutil.which("rsync") and UPLOAD_ROOT.exists():
        try:
            subprocess.run(["rsync", "-a", f"{UPLOAD_ROOT}{os.sep}", str(target)], stderr=subprocess.PIPE, check=True, timeout=3600)
            status = "completed"
            message = "rsync completed"
        except Exception as error:
            status = "failed"
            message = str(error)
    else:
        if UPLOAD_ROOT.exists():
            shutil.copytree(UPLOAD_ROOT, target, dirs_exist_ok=True)
        else:
            target.mkdir(parents=True, exist_ok=True)
        status = "completed"
        message = "rsync unavailable; used shutil copy fallback."
    _prune_old_backups(FILE_BACKUP_ROOT, "files_*")
    return _add_backup_job("files", status, str(target), message)


def full_backup() -> dict[str, Any]:
    return {"database": database_backup(), "files": files_backup()}


def list_backups() -> dict[str, Any]:
    store = load_store()
    return {"jobs": store.get("backup_jobs", []), "database": sorted(str(path) for path in DATABASE_BACKUP_ROOT.glob("backup_*.sql")), "files": sorted(str(path) for path in FILE_BACKUP_ROOT.glob("files_*"))}


def restore_backup(path: str) -> dict[str, Any]:
    before = full_backup()
    backup_path = Path(path)
    if not backup_path.exists():
        raise FileNotFoundError("Backup file not found.")
    database_url = os.environ.get("DATABASE_URL", "").strip()
    if backup_path.suffix == ".sql" and database_url and shutil.which("psql"):
        subprocess.run(["psql", database_url, "-f", str(backup_path)], stderr=subprocess.PIPE, check=True, timeout=3600)
        message = "Restored with psql."
    elif backup_path.suffix == ".sql":
        text = backup_path.read_text(encoding="utf-8", errors="ignore")
        if '"store"' in text:
            try:
                payload = json.loads(text)
                if payload.get("store"):
                    save_store(payload["store"])
                    message = "Restored JSON-compatible store snapshot."
                else:
                    message = "Backup did not contain a local store snapshot."
            except Exception:
                message = "Backup is SQL and psql is unavailable; pre-restore backup was created."
        else:
            message = "Backup is SQL and psql is unavailable; pre-restore backup was created."
    elif backup_path.is_dir():
        shutil.copytree(backup_path, UPLOAD_ROOT, dirs_exist_ok=True)
        message = "Restored upload files from backup directory."
    else:
        message = "Unsupported backup target."
    record_operation("restore_backup", f"Restored {path}: {message}")
    return {"ok": True, "pre_restore_backup": before, "message": message}


def export_records(scope: str = "database", scope_id: str = "", file_type: str = "csv") -> dict[str, Any]:
    ensure_dirs()
    store = load_store()
    if scope == "logs":
        rows = store.get("operation_logs", [])
        filename = f"export_logs_{_timestamp()}.{file_type}"
        path = EXPORT_ROOT / _safe_filename(filename)
        headers = list(rows[0].keys()) if rows else ["id", "user_id", "operation_type", "operation_content", "ip_address", "user_agent", "created_at"]
        with path.open("w", encoding="utf-8-sig", newline="") as handle:
            writer = csv.DictWriter(handle, fieldnames=headers)
            writer.writeheader()
            writer.writerows(rows)
        record_operation("export_logs", f"Exported operation logs to {path}")
        return {"path": str(path), "file_name": path.name, "record_count": len(rows)}
    records = store.get("records", [])
    if scope == "dataset" and scope_id:
        records = [record for record in records if str(record.get("dataset_id")) == str(scope_id)]
    elif scope == "submodule" and scope_id:
        dataset_ids = {str(dataset["id"]) for dataset in list_platform_datasets(scope_id)}
        records = [record for record in records if str(record.get("dataset_id")) in dataset_ids]
    elif scope == "domain" and scope_id:
        sub_ids = {str(sub["id"]) for sub in all_submodules() if sub["domainId"] == scope_id}
        dataset_ids = {str(dataset["id"]) for dataset in store.get("datasets", []) if str(dataset.get("subModuleId")) in sub_ids}
        records = [record for record in records if str(record.get("dataset_id")) in dataset_ids]
    filename = f"export_{scope}_{scope_id or 'all'}_{_timestamp()}.{file_type}"
    path = EXPORT_ROOT / _safe_filename(filename)
    rows = []
    for record in records:
        row = {"id": record.get("id"), "dataset_id": record.get("dataset_id"), **(record.get("system") or {})}
        rows.append(row)
    headers = list(rows[0].keys()) if rows else ["id", "dataset_id", *SYSTEM_FIELDS]
    if file_type == "json":
        path.write_text(json.dumps(rows, ensure_ascii=False, indent=2), encoding="utf-8")
    elif file_type in {"xlsx", "excel"}:
        path = path.with_suffix(".xlsx")
        try:
            import openpyxl  # type: ignore

            workbook = openpyxl.Workbook()
            sheet = workbook.active
            sheet.title = "export"
            sheet.append(headers)
            for row in rows:
                sheet.append([row.get(header, "") for header in headers])
            workbook.save(path)
        except Exception:
            path = path.with_suffix(".csv")
            with path.open("w", encoding="utf-8-sig", newline="") as handle:
                writer = csv.DictWriter(handle, fieldnames=headers)
                writer.writeheader()
                writer.writerows(rows)
    else:
        with path.open("w", encoding="utf-8-sig", newline="") as handle:
            writer = csv.DictWriter(handle, fieldnames=headers)
            writer.writeheader()
            writer.writerows(rows)
    record_operation("export_data", f"Exported {scope}:{scope_id} to {path}")
    return {"path": str(path), "file_name": path.name, "record_count": len(rows)}


def operation_logs(user_id: str = "", operation_type: str = "", start: str = "", end: str = "") -> dict[str, Any]:
    logs = load_store().get("operation_logs", [])
    if user_id:
        logs = [item for item in logs if str(item.get("user_id")) == str(user_id)]
    if operation_type:
        logs = [item for item in logs if str(item.get("operation_type")) == str(operation_type)]
    if start:
        logs = [item for item in logs if str(item.get("created_at", "")) >= start]
    if end:
        logs = [item for item in logs if str(item.get("created_at", "")) <= end]
    return {"logs": logs}


def update_record(record_id: str | int, payload: dict[str, Any]) -> dict[str, Any]:
    store = load_store()
    for record in store.get("records", []):
        if str(record.get("id")) == str(record_id):
            system = record.setdefault("system", {})
            raw = record.setdefault("raw", {})
            updates = payload.get("system") if isinstance(payload.get("system"), dict) else payload
            for key, value in updates.items():
                if key in SYSTEM_FIELDS:
                    system[key] = value
                elif key.startswith("field_"):
                    record[key] = value
                else:
                    raw[key] = value
            record["updated_at"] = _now_iso()
            save_store(store)
            record_operation("record_update", f"Updated record {record_id}")
            return record
    raise KeyError("Record not found.")


def import_backup_file(file_bytes: bytes, filename: str, sub_module_id: str | int = "repository-all-literature-search") -> dict[str, Any]:
    if filename.lower().endswith(".sql"):
        target = DATABASE_BACKUP_ROOT / _safe_filename(filename)
        target.write_bytes(file_bytes)
        return restore_backup(str(target))
    return {"dataset": upload_dataset_file(sub_module_id, file_bytes, filename, {}, Path(filename).stem)}


def scheduler_tick() -> dict[str, Any]:
    store = load_store()
    state = store.setdefault("scheduler_state", {})
    now = datetime.now()
    today = now.strftime("%Y%m%d")
    results = {}
    if now.hour >= 2 and state.get("database_backup_date") != today:
        results["database"] = database_backup(manual=False)
        state["database_backup_date"] = today
    if now.hour >= 3 and state.get("files_backup_date") != today:
        results["files"] = files_backup(manual=False)
        state["files_backup_date"] = today
    save_store(store)
    return results


def start_backup_scheduler() -> None:
    global SCHEDULER_STARTED
    if SCHEDULER_STARTED:
        return
    SCHEDULER_STARTED = True

    def run() -> None:
        while True:
            try:
                scheduler_tick()
            except Exception as error:
                try:
                    record_operation("backup_scheduler_error", str(error))
                except Exception:
                    pass
            time.sleep(60 * 10)

    thread = threading.Thread(target=run, name="platform-backup-scheduler", daemon=True)
    thread.start()
